from __future__ import annotations

import io
import json

from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.core.database import get_db
from app.core.deps import get_current_user
from app.models import Analysis, User
from app.services.analysis_service import analyze_resume

router = APIRouter(prefix="", tags=["analysis"])

_ALLOWED_EXTENSIONS = {".pdf", ".docx"}
_MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB


@router.post("/analyze")
async def analyze(
    file: UploadFile = File(...),
    job_description: str = Form(""),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> dict:
    filename = file.filename or ""
    ext = ("." + filename.lower().rsplit(".", 1)[-1]) if "." in filename else ""
    if ext not in _ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    content = await file.read()
    if len(content) > _MAX_FILE_BYTES:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit.")

    try:
        result = analyze_resume(filename, content, job_description)
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc))

    record = Analysis(
        user_id=current_user.id,
        filename=filename,
        overall_score=result["overall_score"],
        result_json=json.dumps(result),
    )
    db.add(record)
    db.commit()
    db.refresh(record)

    result["id"] = record.id
    result["created_at"] = record.created_at.isoformat() if record.created_at else None
    return result


@router.get("/analyses")
def list_analyses(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> list[dict]:
    records = (
        db.query(Analysis)
        .filter(Analysis.user_id == current_user.id)
        .order_by(Analysis.id.desc())
        .limit(50)
        .all()
    )
    return [r.to_dict() for r in records]


@router.get("/export/analysis/{analysis_id}/docx")
def export_docx(
    analysis_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
) -> StreamingResponse:
    record = db.query(Analysis).filter(
        Analysis.id == analysis_id,
        Analysis.user_id == current_user.id,
    ).first()
    if not record:
        raise HTTPException(status_code=404, detail="Analysis not found.")

    data = json.loads(record.result_json)
    doc = Document()
    doc.add_heading("Resume Analysis Report", 0)
    doc.add_paragraph(f"File: {data.get('filename', '')}")
    doc.add_paragraph(f"Overall Score: {data.get('overall_score', 0)}/100")

    ats = data.get("ats_analysis", {})
    doc.add_heading("ATS Analysis", level=1)
    doc.add_paragraph(f"ATS Score: {ats.get('ats_score', 0)}/100")
    doc.add_paragraph(f"Content Score: {ats.get('content_score', 0)}/100")
    doc.add_paragraph(f"Impact Score: {ats.get('impact_score', 0)}/100")
    doc.add_paragraph(f"Grammar Score: {ats.get('grammar_score', 0)}/100")

    if ats.get("strengths"):
        doc.add_heading("Strengths", level=2)
        for s in ats["strengths"]:
            doc.add_paragraph(f"• {s}")

    if ats.get("issues"):
        doc.add_heading("Issues", level=2)
        for i in ats["issues"]:
            doc.add_paragraph(f"• {i}")

    if data.get("suggestions"):
        doc.add_heading("Suggested Fixes", level=1)
        for s in data["suggestions"]:
            doc.add_paragraph(f"• {s}")

    jd = data.get("jd_match")
    if jd:
        doc.add_heading("Job Description Match", level=1)
        doc.add_paragraph(f"Match Score: {jd.get('match_score', 0)}/100")
        doc.add_paragraph(f"Matched Keywords: {', '.join(jd.get('matched_keywords', []))}")
        doc.add_paragraph(f"Missing Keywords: {', '.join(jd.get('missing_keywords', []))}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=analysis_{analysis_id}.docx"},
    )
